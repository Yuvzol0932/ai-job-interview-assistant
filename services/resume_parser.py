"""简历解析服务：把粘贴文本 / PDF / Word 解析为统一结构。"""

import io

import pdfplumber
from docx import Document

from models.resume import ResumeData

MAX_CHARS = 20000


class ResumeParseError(Exception):
    """简历解析失败的统一异常，信息对用户可读。"""


def _normalize(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(line for line in lines if line)
    return text[:MAX_CHARS]


def parse_resume_text(text: str) -> ResumeData:
    """解析粘贴的纯文本。"""
    return ResumeData(
        content=_normalize(text),
        source_type="paste",
        filename=None,
    )


def parse_resume_file(filename: str, data: bytes) -> ResumeData:
    """解析上传的 PDF / Word 文件。"""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return _parse_pdf(filename, data)
    if ext == "docx":
        return _parse_docx(filename, data)
    if ext == "doc":
        raise ResumeParseError("暂不支持老版 .doc 格式，请另存为 .docx 后再上传，或直接粘贴文本。")
    raise ResumeParseError("仅支持 PDF 或 Word（.docx）文件，请重新上传。")


def _parse_pdf(filename: str, data: bytes) -> ResumeData:
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
    except Exception as exc:
        raise ResumeParseError("PDF 解析失败，文件可能已损坏或为扫描图片，请改用粘贴文本。") from exc
    return ResumeData(
        content=_normalize("\n".join(pages)),
        source_type="pdf",
        filename=filename,
    )


def _parse_docx(filename: str, data: bytes) -> ResumeData:
    try:
        document = Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    except Exception as exc:
        raise ResumeParseError("Word 文件解析失败，文件可能已损坏，请改用粘贴文本。") from exc
    return ResumeData(
        content=_normalize("\n".join(parts)),
        source_type="docx",
        filename=filename,
    )
