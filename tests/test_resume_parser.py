import io

import pytest
from docx import Document

from services.resume_parser import ResumeParseError, parse_resume_file, parse_resume_text


def _build_pdf(text: str) -> bytes:
    content = f"BT /F1 24 Tf 100 700 Td ({text}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length "
        + str(len(content)).encode()
        + b" >>\nstream\n"
        + content
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{index} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF".encode()
    )
    return bytes(out)


def test_parse_text():
    resume = parse_resume_text("  姓名：张三\n\n  学校：XX大学  ")
    assert resume.source_type == "paste"
    assert resume.is_empty is False
    assert "姓名：张三" in resume.content
    assert "\n" in resume.content


def test_parse_empty_text():
    resume = parse_resume_text("   \n  ")
    assert resume.is_empty is True


def test_parse_short_text():
    resume = parse_resume_text("张三")
    assert resume.too_short is True


def test_parse_docx():
    doc = Document()
    doc.add_paragraph("教育经历：XX大学")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "技能"
    table.rows[0].cells[1].text = "Python"
    buffer = io.BytesIO()
    doc.save(buffer)

    resume = parse_resume_file("resume.docx", buffer.getvalue())
    assert resume.source_type == "docx"
    assert "教育经历" in resume.content
    assert "Python" in resume.content


def test_parse_pdf():
    resume = parse_resume_file("resume.pdf", _build_pdf("Hello Resume"))
    assert resume.source_type == "pdf"
    assert "Hello Resume" in resume.content


def test_unsupported_extension():
    with pytest.raises(ResumeParseError):
        parse_resume_file("resume.txt", b"hello")


def test_doc_extension_hint():
    with pytest.raises(ResumeParseError, match="docx"):
        parse_resume_file("resume.doc", b"hello")
